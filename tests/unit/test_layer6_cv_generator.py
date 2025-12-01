"""
Unit tests for Layer 6 CV Generator (Phase 8.2).

Tests STAR-driven CV generation with:
- Competency mix analysis (delivery/process/architecture/leadership)
- STAR scoring and ranking algorithms
- Gap detection (job requirements vs STAR library)
- Hallucination QA pass (verify no invented employers/dates/degrees)
- cv_reasoning generation (gap mitigation strategy)
"""

import pytest
from unittest.mock import MagicMock, patch
from pydantic import ValidationError

pytest.skip("Legacy docx CV generator replaced by markdown flow; skipping legacy assertions.", allow_module_level=True)

from src.common.state import JobState, STARRecord


# ===== FIXTURES =====

@pytest.fixture
def sample_job_state():
    """Sample JobState with all required fields for CV generation."""
    return {
        "job_id": "test_job_123",
        "title": "Senior Engineering Manager",
        "company": "TechCorp",
        "job_description": """
We are seeking a Senior Engineering Manager to lead a team of 8-12 engineers.

Key Responsibilities:
- Lead agile delivery of core platform features with 2-week sprint cycles
- Establish engineering processes (CI/CD, code review, testing standards)
- Design system architecture for multi-tenant SaaS platform
- Mentor junior engineers and build high-performing teams

Requirements:
- 8+ years software engineering experience
- 3+ years engineering management experience
- Track record of scaling teams and systems
- Strong technical background in distributed systems
- Experience with AWS, Kubernetes, microservices architecture
""",
        "pain_points": [
            "Team velocity has decreased 40% over past 6 months",
            "No standardized engineering processes causing quality issues",
            "System architecture cannot scale beyond current 10K users"
        ],
        "strategic_needs": [
            "Build engineering culture of quality and velocity",
            "Establish clear career progression framework for engineers"
        ],
        "selected_stars": [
            {
                "id": "star_001",
                "company": "AdTech Co",
                "role": "Engineering Manager",
                "period": "2020-2023",
                "domain_areas": "Team Leadership, Process Improvement",
                "situation": "Team of 10 engineers with low morale and 3-month release cycles",
                "task": "Transform team culture and accelerate delivery",
                "actions": "Implemented agile practices, established code review standards, mentored 5 engineers to senior roles",
                "results": "Reduced release cycle from 12 weeks to 2 weeks, improved team satisfaction by 60%",
                "metrics": "6x faster releases, 60% higher team engagement, 0 attrition over 2 years",
                "keywords": "agile, team building, process improvement, mentorship"
            },
            {
                "id": "star_002",
                "company": "FinTech Startup",
                "role": "Tech Lead",
                "period": "2018-2020",
                "domain_areas": "System Architecture, Scaling",
                "situation": "Monolithic architecture causing frequent outages during traffic spikes",
                "task": "Redesign system to handle 100x user growth",
                "actions": "Architected microservices migration, built autoscaling infrastructure on AWS/K8s",
                "results": "Scaled from 5K to 500K users with zero downtime",
                "metrics": "100x user growth, 99.99% uptime, $1M cost savings via autoscaling",
                "keywords": "microservices, AWS, Kubernetes, scaling, architecture"
            },
            {
                "id": "star_003",
                "company": "Consulting Firm",
                "role": "Senior Software Engineer",
                "period": "2015-2018",
                "domain_areas": "Software Development, Client Delivery",
                "situation": "Multiple client projects with tight deadlines",
                "task": "Deliver high-quality software solutions on time",
                "actions": "Built 15+ production applications using Java/Python/React, led code reviews",
                "results": "100% on-time delivery, promoted to technical lead",
                "metrics": "15 projects delivered, 100% client satisfaction, promoted in 18 months",
                "keywords": "software development, Java, Python, React, delivery"
            }
        ],
        "candidate_profile": """
Name: Taimoor Alam
Email: taimooralam@example.com
Phone: +1-555-123-4567
LinkedIn: https://linkedin.com/in/taimooralam

EXPERIENCE:
Engineering Manager | AdTech Co | 2020-2023
Tech Lead | FinTech Startup | 2018-2020
Senior Software Engineer | Consulting Firm | 2015-2018

EDUCATION:
B.S. Computer Science | State University | 2015
"""
    }


@pytest.fixture
def all_stars():
    """Complete STAR library for gap detection testing."""
    return [
        {
            "id": "star_001",
            "company": "AdTech Co",
            "role": "Engineering Manager",
            "keywords": "agile, team building, process improvement, mentorship"
        },
        {
            "id": "star_002",
            "company": "FinTech Startup",
            "role": "Tech Lead",
            "keywords": "microservices, AWS, Kubernetes, scaling, architecture"
        },
        {
            "id": "star_003",
            "company": "Consulting Firm",
            "role": "Senior Software Engineer",
            "keywords": "software development, Java, Python, React, delivery"
        },
        {
            "id": "star_004",
            "company": "E-Commerce Platform",
            "role": "DevOps Engineer",
            "keywords": "CI/CD, Jenkins, Docker, monitoring, observability"
        }
    ]


# ===== COMPETENCY MIX ANALYSIS TESTS =====

def test_competency_mix_schema_validation():
    """Test CompetencyMixOutput Pydantic schema validates correctly."""
    from src.layer6.cv_generator import CompetencyMixOutput

    # Valid competency mix
    valid_mix = CompetencyMixOutput(
        delivery=30,
        process=25,
        architecture=25,
        leadership=20,
        reasoning="Job emphasizes delivery (30%) with balanced process/architecture needs"
    )
    assert valid_mix.delivery == 30
    assert valid_mix.process == 25
    assert sum([valid_mix.delivery, valid_mix.process, valid_mix.architecture, valid_mix.leadership]) == 100


def test_competency_mix_sum_to_100():
    """Test that competency percentages must sum to exactly 100."""
    from src.layer6.cv_generator import CompetencyMixOutput

    with pytest.raises(ValidationError):
        CompetencyMixOutput(
            delivery=30,
            process=25,
            architecture=25,
            leadership=25,  # Sums to 105
            reasoning="Invalid sum"
        )


def test_competency_mix_analysis_mock_llm():
    """Test _analyze_competency_mix with mocked LLM response."""
    from src.layer6.cv_generator import CVGenerator

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        # Mock LLM response
        mock_response = MagicMock()
        mock_response.content = """```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Job requires strong delivery focus with balanced process and architecture skills"
}
```"""
        mock_llm.return_value.invoke.return_value = mock_response

        generator = CVGenerator()
        result = generator._analyze_competency_mix("Test job description")

        assert result.delivery == 30
        assert result.leadership == 20


# ===== STAR SCORING AND RANKING TESTS =====

def test_star_scoring_algorithm(sample_job_state, all_stars):
    """Test STAR scoring based on competency alignment."""
    from src.layer6.cv_generator import CVGenerator

    generator = CVGenerator()

    # Mock competency mix emphasizing architecture (40%)
    competency_mix = {
        "delivery": 20,
        "process": 20,
        "architecture": 40,
        "leadership": 20
    }

    # star_002 (architecture/scaling) should score higher than star_001 (team building)
    scores = generator._score_stars(
        all_stars=all_stars,
        competency_mix=competency_mix,
        job_keywords=["microservices", "AWS", "architecture", "scaling"]
    )

    assert scores["star_002"] > scores["star_001"]


def test_star_ranking_selects_top_stars(all_stars):
    """Test that ranking selects top N STARs by score."""
    from src.layer6.cv_generator import CVGenerator

    generator = CVGenerator()

    # Assign mock scores
    scores = {
        "star_001": 85,
        "star_002": 95,
        "star_003": 70,
        "star_004": 80
    }

    top_stars = generator._rank_stars(scores, top_n=2)

    assert len(top_stars) == 2
    assert top_stars[0]["id"] == "star_002"  # Highest score
    assert top_stars[1]["id"] == "star_001"  # Second highest


# ===== GAP DETECTION TESTS =====

def test_gap_detection_identifies_missing_skills(sample_job_state, all_stars):
    """Test gap detection finds required skills not covered by selected STARs."""
    from src.layer6.cv_generator import CVGenerator

    generator = CVGenerator()

    # Job requires: AWS, Kubernetes, CI/CD, distributed systems
    # Selected STARs cover: AWS, Kubernetes (star_002)
    # Missing: CI/CD (only in star_004, not selected)

    gaps = generator._detect_gaps(
        job_description=sample_job_state["job_description"],
        selected_stars=sample_job_state["selected_stars"],
        all_stars=all_stars
    )

    assert len(gaps) > 0
    assert any("CI/CD" in gap or "testing" in gap.lower() for gap in gaps)


def test_gap_detection_no_gaps_when_fully_covered(all_stars):
    """Test that no gaps are detected when all requirements are covered."""
    from src.layer6.cv_generator import CVGenerator

    generator = CVGenerator()

    # Job only requires skills present in selected STARs
    job_description = "We need someone with AWS and Kubernetes experience for microservices architecture"

    selected_stars = [star for star in all_stars if star["id"] in ["star_002"]]

    gaps = generator._detect_gaps(
        job_description=job_description,
        selected_stars=selected_stars,
        all_stars=all_stars
    )

    # Should detect minimal or no gaps since star_002 covers AWS/K8s/microservices
    assert len(gaps) <= 1


# ===== HALLUCINATION QA TESTS =====

def test_hallucination_qa_detects_fake_employer():
    """Test QA pass detects invented employer names."""
    from src.layer6.cv_generator import CVGenerator
    from tenacity import RetryError

    generator = CVGenerator()

    candidate_profile = """
    EXPERIENCE:
    Engineering Manager | AdTech Co | 2020-2023
    Tech Lead | FinTech Startup | 2018-2020
    """

    # CV mentions employer not in candidate profile
    cv_content = """
    EXPERIENCE
    Engineering Manager | FakeCompany Inc | 2020-2023
    Tech Lead | FinTech Startup | 2018-2020
    """

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_response = MagicMock()
        mock_response.content = """```json
{
    "is_valid": false,
    "issues": ["Employer 'FakeCompany Inc' not found in candidate profile"],
    "fabricated_employers": ["FakeCompany Inc"],
    "fabricated_dates": [],
    "fabricated_degrees": []
}
```"""
        mock_llm.return_value.invoke.return_value = mock_response

        # Should raise RetryError after 3 attempts (since QA fails and raises ValueError)
        with pytest.raises(RetryError):
            generator._run_hallucination_qa(cv_content, candidate_profile)


def test_hallucination_qa_detects_fake_dates():
    """Test QA pass detects invented employment dates."""
    from src.layer6.cv_generator import CVGenerator
    from tenacity import RetryError

    generator = CVGenerator()

    candidate_profile = """
    EXPERIENCE:
    Engineering Manager | AdTech Co | 2020-2023
    """

    # CV has wrong dates
    cv_content = """
    EXPERIENCE
    Engineering Manager | AdTech Co | 2018-2023
    """

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_response = MagicMock()
        mock_response.content = """```json
{
    "is_valid": false,
    "issues": ["Date mismatch for AdTech Co: profile says 2020-2023, CV says 2018-2023"],
    "fabricated_employers": [],
    "fabricated_dates": ["2018-2023 for AdTech Co"],
    "fabricated_degrees": []
}
```"""
        mock_llm.return_value.invoke.return_value = mock_response

        # Should raise RetryError after 3 attempts (since QA fails and raises ValueError)
        with pytest.raises(RetryError):
            generator._run_hallucination_qa(cv_content, candidate_profile)


def test_hallucination_qa_passes_valid_cv():
    """Test QA pass approves CV with no fabrications."""
    from src.layer6.cv_generator import CVGenerator

    generator = CVGenerator()

    candidate_profile = """
    EXPERIENCE:
    Engineering Manager | AdTech Co | 2020-2023

    EDUCATION:
    B.S. Computer Science | State University | 2015
    """

    cv_content = """
    EXPERIENCE
    Engineering Manager | AdTech Co | 2020-2023

    EDUCATION
    B.S. Computer Science | State University | 2015
    """

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_response = MagicMock()
        mock_response.content = """```json
{
    "is_valid": true,
    "issues": [],
    "fabricated_employers": [],
    "fabricated_dates": [],
    "fabricated_degrees": []
}
```"""
        mock_llm.return_value.invoke.return_value = mock_response

        result = generator._run_hallucination_qa(cv_content, candidate_profile)

        assert result.is_valid is True
        assert len(result.issues) == 0


# ===== CV REASONING TESTS =====

def test_cv_reasoning_generation(sample_job_state):
    """Test cv_reasoning documents STAR selection and gap mitigation."""
    from src.layer6.cv_generator import CVGenerator

    generator = CVGenerator()

    competency_mix = {
        "delivery": 30,
        "process": 25,
        "architecture": 25,
        "leadership": 20
    }

    gaps = ["CI/CD pipeline automation", "Testing frameworks"]

    reasoning = generator._generate_cv_reasoning(
        competency_mix=competency_mix,
        selected_stars=sample_job_state["selected_stars"],
        gaps=gaps,
        job_title=sample_job_state["title"]
    )

    # Reasoning should mention competency dimensions
    assert "delivery" in reasoning.lower() or "process" in reasoning.lower()

    # Should reference selected STARs
    assert "star" in reasoning.lower() or "engineering manager" in reasoning.lower()

    # Should address gaps
    assert "CI/CD" in reasoning or "gap" in reasoning.lower()


# ===== INTEGRATION TESTS =====

def test_generate_cv_full_pipeline(sample_job_state):
    """Integration test for full CV generation pipeline."""
    from src.layer6.cv_generator import CVGenerator

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        # Create a mock that returns appropriate responses based on call count
        call_count = [0]

        def mock_invoke(messages):
            call_count[0] += 1
            mock_response = MagicMock()

            # First call: competency analysis
            if call_count[0] == 1:
                mock_response.content = """```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Balanced engineering manager role requiring delivery execution and team leadership"
}
```"""
            # Subsequent calls: hallucination QA (passes)
            else:
                mock_response.content = """```json
{
    "is_valid": true,
    "issues": [],
    "fabricated_employers": [],
    "fabricated_dates": [],
    "fabricated_degrees": []
}
```"""
            return mock_response

        mock_llm.return_value.invoke = mock_invoke

        generator = CVGenerator()
        cv_path, cv_reasoning = generator.generate_cv(sample_job_state)

        # Should return path to CV file
        assert cv_path is not None
        assert "TechCorp" in cv_path

        # Should generate reasoning
        assert cv_reasoning is not None
        assert len(cv_reasoning) > 50


def test_generate_cv_retries_on_hallucination_failure(sample_job_state):
    """Test that CV generation retries when hallucination QA fails."""
    from src.layer6.cv_generator import CVGenerator
    from tenacity import RetryError

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        # Mock that always returns QA failure
        qa_fail_response = MagicMock()
        qa_fail_response.content = """```json
{
    "is_valid": false,
    "issues": ["Fake employer detected"],
    "fabricated_employers": ["FakeCorp"],
    "fabricated_dates": [],
    "fabricated_degrees": []
}
```"""

        mock_llm.return_value.invoke.return_value = qa_fail_response

        generator = CVGenerator()

        # Should raise RetryError after exhausting all 3 retry attempts
        with pytest.raises(RetryError):
            generator._validate_cv_content("fake content", sample_job_state["candidate_profile"])


# ===== QUALITY GATE TESTS =====

@pytest.mark.quality_gate
def test_quality_gate_cv_uses_real_employers_only(sample_job_state):
    """
    Quality gate: CV must only use employers from candidate profile.

    Phase 8.2.4: Test that QA rejects CVs with fabricated employers before saving.
    """
    from src.layer6.cv_generator import CVGenerator
    from unittest.mock import patch, MagicMock
    from tenacity import RetryError
    import os

    # Add a fake employer to STARs that's not in candidate_profile
    test_state = sample_job_state.copy()
    fake_star = {
        "id": "star_fake",
        "company": "FakeCorp Industries",  # Not in candidate profile!
        "role": "Fake Role",
        "period": "2025-2026",  # Future dates!
        "domain_areas": "Fabrication",
        "situation": "Fake situation",
        "task": "Fake task",
        "actions": "Fake actions",
        "results": "Fake results with 999% improvement",
        "metrics": "999% fake metrics",
        "keywords": "fake, fabricated, invented"
    }

    test_state["selected_stars"] = sample_job_state["selected_stars"] + [fake_star]
    test_state["all_stars"] = sample_job_state["selected_stars"] + [fake_star]

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_llm_instance = MagicMock()
        mock_llm.return_value = mock_llm_instance

        # Mock competency analysis to succeed
        competency_response = MagicMock()
        competency_response.content = '''```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Job emphasizes delivery, process, architecture, and leadership"
}
```'''

        # Mock hallucination QA to FAIL (detecting the fake employer)
        qa_fail_response = MagicMock()
        qa_fail_response.content = '''```json
{
    "is_valid": false,
    "issues": ["Employer 'FakeCorp Industries' not found in candidate profile", "Date '2025-2026' is in the future"],
    "fabricated_employers": ["FakeCorp Industries"],
    "fabricated_dates": ["2025-2026"],
    "fabricated_degrees": []
}
```'''

        # First call: competency analysis succeeds
        # Subsequent calls: QA fails repeatedly (exhausts retries)
        mock_llm_instance.invoke.side_effect = [competency_response] + [qa_fail_response] * 5

        generator = CVGenerator()

        # Generate CV path to check if file is created
        company_clean = test_state["company"].replace(" ", "_").replace("/", "_")
        title_clean = test_state["title"].replace(" ", "_").replace("/", "_")
        expected_cv_path = f"applications/{company_clean}/{title_clean}/CV_{company_clean}.docx"

        # Clean up any existing file from previous test runs
        if os.path.exists(expected_cv_path):
            os.remove(expected_cv_path)

        # Should raise RetryError because QA fails
        with pytest.raises(RetryError):
            generator.generate_cv(test_state)

        # CRITICAL: Verify that no .docx file was saved (QA prevented saving)
        assert not os.path.exists(expected_cv_path), \
            f"CV file should NOT have been saved when QA detects fabricated employer, but found {expected_cv_path}"

        # Clean up
        if os.path.exists(expected_cv_path):
            os.remove(expected_cv_path)


@pytest.mark.quality_gate
def test_quality_gate_cv_addresses_top_competencies(sample_job_state):
    """Quality gate: CV must highlight STARs aligned with job's top 2 competency dimensions."""
    from src.layer6.cv_generator import CVGenerator

    # This test will be implemented after the generator is built
    # It should verify that the CV emphasizes the right competencies
    pass


# ===== PHASE 8.2 ROADMAP TESTS =====

def test_cv_generator_uses_all_stars_not_selected_stars(sample_job_state):
    """
    Phase 8.2.1: CV generator should use all_stars (full library) for scoring/ranking,
    not just the pre-selected selected_stars.

    ROADMAP requirement: "Use full STAR library and ranking in generate_cv"
    """
    from src.layer6.cv_generator import CVGenerator
    from unittest.mock import patch, MagicMock

    # Add more STARs to all_stars than selected_stars
    all_stars = sample_job_state["selected_stars"] + [
        {
            "id": "star_004",
            "company": "Enterprise Co",
            "role": "Senior Architect",
            "period": "2013-2015",
            "domain_areas": "Cloud Architecture, Cost Optimization",
            "situation": "High cloud costs threatening profitability",
            "task": "Reduce AWS spending by 50% without performance impact",
            "actions": "Implemented autoscaling, reserved instances, and spot fleets",
            "results": "Reduced monthly AWS costs from $200K to $75K",
            "metrics": "$1.5M annual savings, maintained 99.95% uptime",
            "keywords": "AWS, cost optimization, cloud architecture, autoscaling"
        },
        {
            "id": "star_005",
            "company": "Startup Inc",
            "role": "CTO",
            "period": "2011-2013",
            "domain_areas": "Leadership, Hiring, Technical Strategy",
            "situation": "Early-stage startup needing technical foundation",
            "task": "Build engineering team from scratch and define tech stack",
            "actions": "Hired 12 engineers, established CI/CD pipeline, chose microservices architecture",
            "results": "Delivered MVP in 6 months, scaled to Series A",
            "metrics": "12 hires in 9 months, 0 production outages, raised $5M Series A",
            "keywords": "hiring, team building, startup, CTO, technical leadership"
        }
    ]

    # Update state with all_stars and updated candidate profile
    test_state = sample_job_state.copy()
    test_state["all_stars"] = all_stars

    # Add the new companies to candidate profile so hallucination QA passes
    test_state["candidate_profile"] = sample_job_state["candidate_profile"] + """

Enterprise Co | Senior Architect | 2013-2015
- Reduced AWS spending by 50% without performance impact
- Implemented autoscaling, reserved instances, and spot fleets

Startup Inc | CTO | 2011-2013
- Built engineering team from scratch and defined tech stack
- Delivered MVP in 6 months, scaled to Series A
"""

    # Mock ChatOpenAI before creating generator
    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        # Mock competency mix response
        mock_llm_instance = MagicMock()
        mock_llm.return_value = mock_llm_instance

        competency_response = MagicMock()
        competency_response.content = '''```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Job emphasizes delivery (30%), process (25%), architecture (25%), and leadership (20%)"
}
```'''

        mock_llm_instance.invoke.return_value = competency_response

        # Create generator with mocked LLM
        generator = CVGenerator()

        # Mock hallucination QA to always pass (focus of this test is STAR scoring/ranking)
        from src.layer6.cv_generator import HallucinationQAOutput
        mock_qa_result = HallucinationQAOutput(
            is_valid=True,
            issues=[],
            fabricated_employers=[],
            fabricated_dates=[],
            fabricated_degrees=[]
        )

        # Spy on _score_stars and _rank_stars methods
        with patch.object(generator, '_score_stars', wraps=generator._score_stars) as mock_score:
            with patch.object(generator, '_rank_stars', wraps=generator._rank_stars) as mock_rank:
                with patch.object(generator, '_validate_cv_content', return_value=mock_qa_result):
                    cv_path, cv_reasoning = generator.generate_cv(test_state)

                    # ROADMAP 8.2 verification: _score_stars should be called with all_stars (5 STARs)
                    mock_score.assert_called_once()
                    call_args = mock_score.call_args
                    stars_argument = call_args[0][0]  # First positional arg: all_stars

                    # Should have scored ALL 5 stars, not just the 3 selected_stars
                    assert len(stars_argument) == 5, f"Expected _score_stars to receive 5 STARs from all_stars, got {len(stars_argument)}"

                    # Verify it includes stars not in selected_stars
                    star_ids = [s["id"] for s in stars_argument]
                    assert "star_004" in star_ids, "Should include star_004 from all_stars"
                    assert "star_005" in star_ids, "Should include star_005 from all_stars"

                    # _rank_stars should also be called
                    mock_rank.assert_called_once()


def test_cv_generator_creates_docx_with_proper_structure(sample_job_state):
    """
    Phase 8.2.2: CV generator should create .docx file (not .txt) with proper structure.

    ROADMAP requirement: "assert a .docx filename and basic structural markers"
    """
    from src.layer6.cv_generator import CVGenerator
    from unittest.mock import patch, MagicMock
    from docx import Document
    import os

    # Update state with all_stars
    test_state = sample_job_state.copy()
    test_state["all_stars"] = sample_job_state["selected_stars"]

    # Mock ChatOpenAI before creating generator
    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_llm_instance = MagicMock()
        mock_llm.return_value = mock_llm_instance

        competency_response = MagicMock()
        competency_response.content = '''```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Job emphasizes delivery, process, architecture, and leadership"
}
```'''

        mock_llm_instance.invoke.return_value = competency_response

        generator = CVGenerator()

        # Mock hallucination QA to pass
        from src.layer6.cv_generator import HallucinationQAOutput
        mock_qa = HallucinationQAOutput(
            is_valid=True,
            issues=[],
            fabricated_employers=[],
            fabricated_dates=[],
            fabricated_degrees=[]
        )

        with patch.object(generator, '_validate_cv_content', return_value=mock_qa):
            cv_path, cv_reasoning = generator.generate_cv(test_state)

            # ROADMAP 8.2.2 verification: Should create .docx file (not .txt)
            assert cv_path.endswith(".docx"), f"Expected .docx file, got {cv_path}"

            # Verify file exists
            assert os.path.exists(cv_path), f"CV file not found at {cv_path}"

            # Load the .docx and verify structure
            doc = Document(cv_path)

            # Extract all text from the document
            all_text = "\n".join([para.text for para in doc.paragraphs])

            # ROADMAP requirement: Check for required section headings
            assert "Professional Summary" in all_text, "Missing 'Professional Summary' heading"
            assert "Key Achievements" in all_text, "Missing 'Key Achievements' heading"
            assert "Professional Experience" in all_text, "Missing 'Professional Experience' heading"
            assert "Education" in all_text or "Certifications" in all_text, "Missing 'Education & Certifications' heading"

            # Verify candidate name appears
            assert "Taimoor Alam" in all_text, "Candidate name should appear in CV"

            # Verify at least one company from STARs appears
            companies = [star["company"] for star in sample_job_state["selected_stars"]]
            assert any(company in all_text for company in companies), "At least one STAR company should appear"

            # Clean up
            if os.path.exists(cv_path):
                os.remove(cv_path)


# ===== PHASE 8.2.3: EDGE CASE TESTS =====

def test_cv_generator_handles_empty_star_list(sample_job_state):
    """
    Phase 8.2.3: CV generator should handle empty STAR list gracefully.

    ROADMAP requirement: "Emit a minimal but valid CV derived from candidate_profile only"
    """
    from src.layer6.cv_generator import CVGenerator
    from unittest.mock import patch, MagicMock
    import os

    # Empty STAR list
    test_state = sample_job_state.copy()
    test_state["selected_stars"] = []
    test_state["all_stars"] = []

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_llm_instance = MagicMock()
        mock_llm.return_value = mock_llm_instance

        competency_response = MagicMock()
        competency_response.content = '''```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Job emphasizes delivery, process, architecture, and leadership"
}
```'''

        mock_llm_instance.invoke.return_value = competency_response

        generator = CVGenerator()

        # Generate CV - should succeed even with empty STAR list
        cv_path, cv_reasoning = generator.generate_cv(test_state)

        # Should generate a CV path (not return None)
        assert cv_path is not None, "Should generate CV even without STARs"
        assert cv_path.endswith(".docx"), "Should create .docx file"

        # cv_reasoning should explain the limitation
        assert "no" in cv_reasoning.lower() or "minimal" in cv_reasoning.lower() or "missing" in cv_reasoning.lower(), \
            "cv_reasoning should explain lack of STAR data"

        # CV file should exist
        assert os.path.exists(cv_path), f"CV file should exist at {cv_path}"

        # Clean up
        if os.path.exists(cv_path):
            os.remove(cv_path)


def test_cv_generator_handles_single_star(sample_job_state):
    """
    Phase 8.2.3: CV generator should handle single STAR gracefully.

    ROADMAP requirement: Handle minimal-STAR scenarios
    """
    from src.layer6.cv_generator import CVGenerator
    from unittest.mock import patch, MagicMock
    from docx import Document
    import os

    # Only one STAR
    test_state = sample_job_state.copy()
    test_state["selected_stars"] = [sample_job_state["selected_stars"][0]]  # Just first STAR
    test_state["all_stars"] = [sample_job_state["selected_stars"][0]]

    with patch('src.layer6.cv_generator.create_tracked_cv_llm') as mock_llm:
        mock_llm_instance = MagicMock()
        mock_llm.return_value = mock_llm_instance

        competency_response = MagicMock()
        competency_response.content = '''```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Job emphasizes delivery, process, architecture, and leadership"
}
```'''

        mock_llm_instance.invoke.return_value = competency_response

        generator = CVGenerator()

        # Mock QA to pass
        from src.layer6.cv_generator import HallucinationQAOutput
        mock_qa = HallucinationQAOutput(
            is_valid=True,
            issues=[],
            fabricated_employers=[],
            fabricated_dates=[],
            fabricated_degrees=[]
        )

        with patch.object(generator, '_validate_cv_content', return_value=mock_qa):
            cv_path, cv_reasoning = generator.generate_cv(test_state)

            # Should succeed
            assert cv_path is not None
            assert cv_path.endswith(".docx")

            # Verify file exists and contains the single STAR's company
            assert os.path.exists(cv_path)

            doc = Document(cv_path)
            all_text = "\n".join([para.text for para in doc.paragraphs])

            # Should contain the single STAR's company
            assert sample_job_state["selected_stars"][0]["company"] in all_text

            # Clean up
            if os.path.exists(cv_path):
                os.remove(cv_path)
