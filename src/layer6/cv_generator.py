"""
Layer 6 CV Generator (Phase 8.2) - STAR-Driven CV Generation.

Implements sophisticated CV tailoring with:
1. Competency mix analysis (delivery/process/architecture/leadership dimensions)
2. STAR scoring and ranking (algorithmic selection based on job fit)
3. Gap detection (identify missing skills vs job requirements)
4. Hallucination QA pass (verify no invented employers/dates/degrees)
5. cv_reasoning generation (document tailoring rationale and gap mitigation)

This goes beyond basic CV templating to strategically select and present
achievements that align with the target role's competency profile.
"""

import logging
import re
import json
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from pydantic import BaseModel, Field, validator
from tenacity import retry, stop_after_attempt, wait_exponential

from src.common.config import Config
from src.common.llm_factory import create_tracked_cv_llm
from src.common.state import JobState, STARRecord
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===== PYDANTIC SCHEMAS =====

class CompetencyMixOutput(BaseModel):
    """
    Competency dimension analysis for a job (Phase 8.2).

    Classifies job requirements into 4 dimensions that sum to 100%.
    Used to algorithmically score and rank STARs by relevance.

    ROADMAP Phase 8.2 Quality Gates:
    - All 4 dimensions must be present
    - Percentages must sum to exactly 100
    - Reasoning must explain the mix based on job description
    """
    delivery: int = Field(..., ge=0, le=100, description="% focus on shipping features/products")
    process: int = Field(..., ge=0, le=100, description="% focus on engineering processes/quality")
    architecture: int = Field(..., ge=0, le=100, description="% focus on system design/technical decisions")
    leadership: int = Field(..., ge=0, le=100, description="% focus on people management/mentorship")
    reasoning: str = Field(..., min_length=50, description="Explanation of competency breakdown")

    @validator('leadership')
    def validate_sum_to_100(cls, v, values):
        """Ensure all competencies sum to exactly 100%."""
        delivery = values.get('delivery', 0)
        process = values.get('process', 0)
        architecture = values.get('architecture', 0)
        total = delivery + process + architecture + v

        if total != 100:
            raise ValueError(
                f"Competencies must sum to 100%, got {total}% "
                f"(delivery={delivery}, process={process}, architecture={architecture}, leadership={v})"
            )

        return v


class HallucinationQAOutput(BaseModel):
    """
    Hallucination QA validation result (Phase 8.2).

    LLM-based quality check that verifies CV contains no fabricated:
    - Employer names not in candidate profile
    - Employment dates that don't match profile
    - Degrees or schools not in education history

    Quality Gate: is_valid must be True before CV is saved.
    """
    is_valid: bool = Field(..., description="True if CV passes all hallucination checks")
    issues: List[str] = Field(default_factory=list, description="List of detected issues")
    fabricated_employers: List[str] = Field(default_factory=list, description="Employer names not in profile")
    fabricated_dates: List[str] = Field(default_factory=list, description="Date mismatches")
    fabricated_degrees: List[str] = Field(default_factory=list, description="Degrees/schools not in profile")


# ===== CV GENERATOR CLASS =====

class CVGenerator:
    """
    STAR-driven CV generator with hallucination prevention (Phase 8.2).

    Pipeline:
    1. Analyze job competency mix (delivery/process/architecture/leadership)
    2. Score all STARs against competency mix and job keywords
    3. Rank STARs and select top N for inclusion
    4. Detect gaps (required skills not covered by selected STARs)
    5. Generate cv_reasoning (explain STAR choices and gap mitigation)
    6. Build CV document with enhanced structure
    7. Run hallucination QA pass to verify no fabrications
    8. Retry if QA fails, otherwise save CV

    Temperature: 0.3 (ANALYTICAL_TEMPERATURE) for competency analysis and QA
    Model: GPT-4o (Config.DEFAULT_MODEL)
    """

    def __init__(self):
        """
        Initialize CV generator with LLM client.

        Selects LLM provider based on Config:
        - Anthropic (direct): Uses ChatAnthropic with Anthropic API key
        - OpenRouter (proxy): Uses ChatOpenAI with OpenRouter base URL
        - OpenAI (direct): Uses ChatOpenAI with OpenAI API key
        """
        # Logger for internal operations
        self.logger = logging.getLogger(__name__)

        # GAP-066: Token tracking enabled via factory
        self.llm = create_tracked_cv_llm(layer="layer6_cv_v2")
        self.logger.info(f"Using tracked CV LLM (provider: {Config.get_cv_llm_provider()})")

    # ===== COMPETENCY MIX ANALYSIS =====

    # V2 Enhanced Prompt - A/B Testing Phase 8.2
    SYSTEM_PROMPT_COMPETENCY_MIX = """You are an expert technical recruiter analyzing job requirements for CV tailoring.

Your task: Classify a job into 4 competency dimensions to select the most relevant candidate achievements.

DIMENSIONS (must sum to exactly 100%):
1. **Delivery** (0-100%): Shipping features, building products, writing code, executing projects
2. **Process** (0-100%): Engineering practices, quality, CI/CD, testing, code review standards
3. **Architecture** (0-100%): System design, technical strategy, infrastructure, scalability decisions
4. **Leadership** (0-100%): People management, mentorship, team building, culture, hiring

ANALYSIS FRAMEWORK:

STEP 1: KEYWORD EXTRACTION
Identify explicit competency signals in the job description:
- Delivery keywords: "build", "ship", "develop", "implement", "deliver"
- Process keywords: "CI/CD", "testing", "quality", "code review", "standards"
- Architecture keywords: "design", "scale", "architecture", "infrastructure", "system"
- Leadership keywords: "lead", "mentor", "manage", "team", "hire", "culture"

STEP 2: CONTEXT ANALYSIS
Beyond keywords, analyze the role's seniority and scope:
- Junior roles: Higher delivery weight
- Staff/Principal: Higher architecture weight
- Team Lead/Manager: Higher leadership weight
- DevOps/Platform: Higher process weight

STEP 3: BALANCE CHECK
Verify the mix makes sense for this specific role before outputting.

CRITICAL RULES:
- Percentages MUST sum to exactly 100
- Reasoning MUST quote specific phrases from the job description
- Each percentage MUST be justified with evidence from the JD

Output ONLY valid JSON:
```json
{
    "delivery": 30,
    "process": 25,
    "architecture": 25,
    "leadership": 20,
    "reasoning": "Evidence: [specific JD quotes for each dimension]"
}
```"""

    USER_PROMPT_COMPETENCY_MIX_TEMPLATE = """Analyze this job description and classify into competency dimensions:

JOB TITLE: {title}
COMPANY: {company}

JOB DESCRIPTION:
{job_description}

Output JSON with delivery/process/architecture/leadership percentages (must sum to 100) and reasoning."""

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    def _analyze_competency_mix(self, job_description: str, title: str = "", company: str = "") -> CompetencyMixOutput:
        """
        Analyze job description to determine competency mix.

        Returns validated CompetencyMixOutput with 4 dimensions summing to 100%.
        Retries up to 3 times on validation failures.
        """
        prompt = self.USER_PROMPT_COMPETENCY_MIX_TEMPLATE.format(
            title=title,
            company=company,
            job_description=job_description
        )

        messages = [
            {"role": "system", "content": self.SYSTEM_PROMPT_COMPETENCY_MIX},
            {"role": "user", "content": prompt}
        ]

        response = self.llm.invoke(messages)
        content = response.content.strip()

        # Extract JSON from markdown code blocks
        json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)

        # Parse and validate
        data = json.loads(content)
        return CompetencyMixOutput(**data)

    # ===== STAR SCORING AND RANKING =====

    def _score_stars(
        self,
        all_stars: List[Dict],
        competency_mix: Dict[str, int],
        job_keywords: List[str]
    ) -> Dict[str, float]:
        """
        Score each STAR based on competency alignment and keyword matches.

        Scoring algorithm:
        - 60% weight: Competency dimension alignment
        - 40% weight: Keyword match score

        Returns dict mapping star_id -> score (0-100).
        """
        scores = {}

        # Extract keywords from job (simplified - could use TF-IDF in production)
        job_keywords_lower = [k.lower() for k in job_keywords]

        for star in all_stars:
            star_id = star.get("id", star.get("company", "unknown"))

            # Keyword match score (0-40 points)
            star_keywords = star.get("keywords", "").lower().split(",")
            star_keywords = [k.strip() for k in star_keywords]

            keyword_matches = sum(1 for sk in star_keywords if any(jk in sk for jk in job_keywords_lower))
            keyword_score = min(40, keyword_matches * 10)

            # Competency alignment score (0-60 points)
            # Map STAR to competency dimensions based on domain_areas and keywords
            star_competencies = self._infer_star_competencies(star)

            competency_score = 0
            for dimension, weight in competency_mix.items():
                # Skip non-numeric fields like 'reasoning'
                if not isinstance(weight, (int, float)):
                    continue
                star_strength = star_competencies.get(dimension, 0)
                competency_score += (star_strength * weight / 100) * 0.6

            competency_score = competency_score * 100  # Scale to 0-60

            total_score = keyword_score + competency_score
            scores[star_id] = total_score

        return scores

    def _infer_star_competencies(self, star: Dict) -> Dict[str, float]:
        """
        Infer STAR's competency strengths (0-1 scale for each dimension).

        Uses keywords and domain_areas to classify STAR into dimensions.
        Returns dict with delivery/process/architecture/leadership scores.
        """
        keywords = star.get("keywords", "").lower()
        domain = star.get("domain_areas", "").lower()
        role = star.get("role", "").lower()

        competencies = {
            "delivery": 0.0,
            "process": 0.0,
            "architecture": 0.0,
            "leadership": 0.0
        }

        # Delivery indicators
        if any(k in keywords or k in domain for k in ["shipped", "built", "developed", "delivery", "implementation"]):
            competencies["delivery"] = 0.8

        # Process indicators
        if any(k in keywords or k in domain for k in ["ci/cd", "testing", "quality", "code review", "agile", "process"]):
            competencies["process"] = 0.8

        # Architecture indicators
        if any(k in keywords or k in domain for k in ["architecture", "system design", "scaling", "infrastructure", "microservices", "aws", "kubernetes"]):
            competencies["architecture"] = 0.8

        # Leadership indicators
        if any(k in keywords or k in domain or k in role for k in ["manager", "lead", "mentorship", "team building", "hiring", "leadership"]):
            competencies["leadership"] = 0.8

        return competencies

    def _rank_stars(self, scores: Dict[str, float], top_n: int = 3) -> List[Dict]:
        """
        Rank STARs by score and return top N.

        Returns list of STAR IDs sorted by score (highest first).
        """
        sorted_stars = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        return [{"id": star_id, "score": score} for star_id, score in sorted_stars[:top_n]]

    # ===== GAP DETECTION =====

    def _detect_gaps(
        self,
        job_description: str,
        selected_stars: List[Dict],
        all_stars: List[Dict]
    ) -> List[str]:
        """
        Detect skill gaps: required skills not covered by selected STARs.

        Algorithm:
        1. Extract key requirements from job description (manual keyword extraction)
        2. Check which requirements are covered by selected STARs
        3. Return list of uncovered requirements
        4. For each gap, check if it's covered by non-selected STARs (could be mentioned in reasoning)

        Returns list of gap descriptions (e.g., "CI/CD pipeline automation not demonstrated").
        """
        gaps = []

        # Common skill patterns to check (could be enhanced with NER/LLM extraction)
        skill_patterns = [
            (r'\bCI/CD\b', "CI/CD pipeline automation"),
            (r'\btesting\b|\btest automation\b', "Automated testing frameworks"),
            (r'\bKubernetes\b|\bK8s\b', "Kubernetes orchestration"),
            (r'\bAWS\b|\bcloud\b', "AWS cloud infrastructure"),
            (r'\bmicroservices\b', "Microservices architecture"),
            (r'\bDocker\b|\bcontainers\b', "Container technologies"),
            (r'\bmonitoring\b|\bobservability\b', "System monitoring and observability"),
            (r'\bagile\b|\bscrum\b', "Agile development practices")
        ]

        # Check each pattern
        for pattern, gap_description in skill_patterns:
            if re.search(pattern, job_description, re.IGNORECASE):
                # Check if covered by selected STARs
                covered = False
                for star in selected_stars:
                    star_text = f"{star.get('keywords', '')} {star.get('domain_areas', '')} {star.get('actions', '')}".lower()
                    if re.search(pattern, star_text, re.IGNORECASE):
                        covered = True
                        break

                if not covered:
                    gaps.append(gap_description)

        return gaps

    # ===== HALLUCINATION QA =====

    # V2 Enhanced Hallucination QA - A/B Testing Phase 8.2
    SYSTEM_PROMPT_HALLUCINATION_QA = """You are a strict quality assurance agent detecting fabricated information in CVs.

Your job: Verify EVERY claim in the CV against the candidate profile (source of truth).

CRITICAL VERIFICATION CHECKLIST:

1. **EMPLOYERS** (Zero tolerance for fabrication)
   - Every company name MUST exist in candidate profile
   - Allow minor variations: "Seven.One Entertainment Group" = "Seven One Entertainment"
   - DO NOT allow: Completely new company names not in profile

2. **EMPLOYMENT DATES** (Zero tolerance for fabrication)
   - Every date range MUST match or be subset of profile dates
   - Allow format variations: "2020–Present" = "2020-Present" = "2020 to Present"
   - DO NOT allow: Wrong years, extended ranges, compressed timelines

3. **METRICS & ACHIEVEMENTS** (Zero tolerance for fabrication)
   - Every metric (%, $, numbers) MUST appear in candidate profile
   - Allow rephrasing: "75% reduction" = "reduced by 75%"
   - DO NOT allow: Made-up percentages, invented dollar amounts, fabricated team sizes

4. **DEGREES & CERTIFICATIONS** (Zero tolerance for fabrication)
   - Every school/degree MUST exist in candidate profile
   - Allow: Abbreviations, formatting differences
   - DO NOT allow: Fabricated schools, fake degrees

5. **METRICS SOURCE CHECK** (NEW)
   - Cross-reference each metric in CV against STAR records or master CV
   - Flag any metric not traceable to source

VERIFICATION PROCESS:
Step 1: List all companies in CV, verify each against profile
Step 2: List all date ranges in CV, verify each against profile
Step 3: List all metrics in CV, verify each against profile or STARs
Step 4: List all education in CV, verify each against profile

Output JSON with is_valid=true ONLY if ALL checks pass.

```json
{
    "is_valid": false,
    "issues": ["[Specific issue with source reference]"],
    "fabricated_employers": ["[Company not in profile]"],
    "fabricated_dates": ["[Mismatch with explanation]"],
    "fabricated_degrees": [],
    "unverifiable_metrics": ["[Metric not found in source]"]
}
```"""

    USER_PROMPT_HALLUCINATION_QA_TEMPLATE = """Check this CV for fabricated information against the candidate profile.

CANDIDATE PROFILE (SOURCE OF TRUTH):
{candidate_profile}

GENERATED CV (TO VERIFY):
{cv_content}

Output JSON with is_valid, issues, fabricated_employers, fabricated_dates, fabricated_degrees."""

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    def _run_hallucination_qa(self, cv_content: str, candidate_profile: str) -> HallucinationQAOutput:
        """
        Run LLM-based hallucination QA to verify CV contains no fabrications.

        Checks:
        - All employers in CV exist in candidate profile
        - All dates match candidate profile
        - All degrees/schools match candidate profile

        Raises ValueError if QA fails (to trigger retry in generate_cv).
        """
        prompt = self.USER_PROMPT_HALLUCINATION_QA_TEMPLATE.format(
            candidate_profile=candidate_profile,
            cv_content=cv_content
        )

        messages = [
            {"role": "system", "content": self.SYSTEM_PROMPT_HALLUCINATION_QA},
            {"role": "user", "content": prompt}
        ]

        response = self.llm.invoke(messages)
        content = response.content.strip()

        # Extract JSON
        json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)

        data = json.loads(content)
        result = HallucinationQAOutput(**data)

        # If QA failed, raise ValueError to trigger retry
        if not result.is_valid:
            error_msg = f"Hallucination QA failed: {', '.join(result.issues)}"
            self.logger.warning(f"{error_msg}")
            if result.fabricated_employers:
                self.logger.warning(f"  Fabricated employers: {result.fabricated_employers}")
            if result.fabricated_dates:
                self.logger.warning(f"  Fabricated dates: {result.fabricated_dates}")
            if result.fabricated_degrees:
                self.logger.warning(f"  Fabricated degrees: {result.fabricated_degrees}")
            raise ValueError(error_msg)

        return result

    def _validate_cv_content(self, cv_content: str, candidate_profile: str) -> HallucinationQAOutput:
        """Wrapper for _run_hallucination_qa that can be mocked in tests."""
        return self._run_hallucination_qa(cv_content, candidate_profile)

    # ===== CV REASONING GENERATION =====

    def _generate_cv_reasoning(
        self,
        competency_mix: Dict[str, int],
        selected_stars: List[Dict],
        gaps: List[str],
        job_title: str
    ) -> str:
        """
        Generate cv_reasoning field documenting STAR selection and gap mitigation.

        Reasoning should explain:
        1. Why these STARs were selected (competency alignment)
        2. How they address top job requirements
        3. What gaps exist and how they're mitigated (or acknowledged)

        Returns 2-4 paragraph reasoning string.
        """
        reasoning_parts = []

        # Part 1: Competency mix analysis
        # Filter out non-numeric fields (like 'reasoning')
        numeric_competencies = {k: v for k, v in competency_mix.items() if isinstance(v, (int, float))}
        top_competencies = sorted(numeric_competencies.items(), key=lambda x: x[1], reverse=True)[:2]
        competency_summary = f"Job analysis shows '{job_title}' requires {top_competencies[0][0]} ({top_competencies[0][1]}%) and {top_competencies[1][0]} ({top_competencies[1][1]}%) as primary competencies."
        reasoning_parts.append(competency_summary)

        # Part 2: STAR selection rationale
        star_ids = [s.get("id", s.get("company", "unknown")) for s in selected_stars]
        star_summary = f"Selected {len(selected_stars)} STARs ({', '.join(star_ids)}) that best demonstrate these competencies with quantified business impact."
        reasoning_parts.append(star_summary)

        # Part 3: Gap mitigation (if gaps exist)
        if gaps:
            gap_summary = f"Identified {len(gaps)} skill gaps: {', '.join(gaps[:3])}{'...' if len(gaps) > 3 else ''}. These are addressed through: (1) highlighting transferable skills from selected STARs, (2) emphasizing learning agility in summary, (3) noting relevant context in candidate profile."
            reasoning_parts.append(gap_summary)
        else:
            reasoning_parts.append("No significant skill gaps detected - selected STARs provide comprehensive coverage of job requirements.")

        return " ".join(reasoning_parts)

    # ===== CV GENERATION =====

    def generate_cv(self, state: JobState) -> Tuple[str, str]:
        """
        Generate STAR-driven tailored CV (Phase 8.2 main entry point).

        Pipeline:
        1. Analyze competency mix
        2. Score and rank all STARs
        3. Detect gaps
        4. Generate cv_reasoning
        5. Build CV document
        6. Run hallucination QA
        7. Save CV if QA passes

        Returns (cv_path, cv_reasoning) tuple.
        """
        self.logger.info("="*60)
        self.logger.info("CV GENERATOR (Phase 8.2)")
        self.logger.info("="*60)

        # Step 1: Competency mix analysis
        self.logger.info("Analyzing job competency mix...")
        competency_mix = self._analyze_competency_mix(
            job_description=state["job_description"],
            title=state["title"],
            company=state["company"]
        )
        self.logger.info(f"Competency mix: Delivery={competency_mix.delivery}%, Process={competency_mix.process}%, Architecture={competency_mix.architecture}%, Leadership={competency_mix.leadership}%")

        # Step 2: Score and rank STARs (Phase 8.2: Use full STAR library)
        self.logger.info("Scoring and ranking STARs...")
        # NOTE: Use `or []` pattern to handle None values (key may exist with None)
        all_stars = state.get("all_stars") or state.get("selected_stars") or []

        # Phase 8.2.3: Handle empty/minimal STAR list gracefully
        if not all_stars:
            self.logger.warning("No STARs available - generating minimal CV from profile only")
            return self._generate_minimal_cv(state, competency_mix)

        # Extract job keywords for scoring
        job_keywords = self._extract_keywords(state["job_description"])

        # Score all STARs against job requirements
        competency_dict = competency_mix.dict()
        star_scores = self._score_stars(all_stars, competency_dict, job_keywords)
        self.logger.info(f"Scored {len(star_scores)} STARs")

        # Rank and select top 3-5 STARs
        top_n = min(5, max(3, len(all_stars)))
        ranked_stars = self._rank_stars(star_scores, top_n=top_n)
        self.logger.info(f"Selected top {len(ranked_stars)} STARs for CV")

        # Get full STAR objects for selected IDs
        star_id_to_record = {star["id"]: star for star in all_stars}
        selected_stars_for_cv = []
        for ranked in ranked_stars:
            star_id = ranked["id"]
            if star_id in star_id_to_record:
                selected_stars_for_cv.append(star_id_to_record[star_id])

        # Fallback: if no stars ranked, use all available
        if not selected_stars_for_cv:
            selected_stars_for_cv = all_stars[:top_n]

        # Step 3: Detect gaps
        self.logger.info("Detecting skill gaps...")
        gaps = self._detect_gaps(
            job_description=state["job_description"],
            selected_stars=selected_stars_for_cv,
            all_stars=all_stars
        )
        if gaps:
            self.logger.info(f"Found {len(gaps)} gaps: {', '.join(gaps[:3])}")
        else:
            self.logger.info("No significant gaps detected")

        # Step 4: Generate cv_reasoning
        self.logger.info("Generating CV reasoning...")
        cv_reasoning = self._generate_cv_reasoning(
            competency_mix=competency_mix.dict(),
            selected_stars=selected_stars_for_cv,
            gaps=gaps,
            job_title=state["title"]
        )

        # Step 5: Build CV document
        self.logger.info("Building CV document...")
        cv_doc, cv_text_content = self._build_cv_document(state, competency_mix, selected_stars_for_cv)

        # Step 6: Run hallucination QA (on text content)
        self.logger.info("Running hallucination QA pass...")
        try:
            qa_result = self._validate_cv_content(cv_text_content, state["candidate_profile"])
            self.logger.info("Hallucination QA passed")
        except ValueError as e:
            self.logger.error(f"Hallucination QA failed: {e}")
            raise

        # Step 7: Save CV (.docx)
        cv_path = self._save_cv_document(state, cv_doc)
        self.logger.info(f"CV saved: {cv_path}")

        return cv_path, cv_reasoning

    def _extract_keywords(self, job_description: str) -> List[str]:
        """Extract key technical keywords from job description."""
        # Simplified keyword extraction (could use TF-IDF or LLM in production)
        common_keywords = [
            "AWS", "Kubernetes", "microservices", "CI/CD", "Docker", "Python", "Java",
            "React", "distributed systems", "agile", "scrum", "leadership", "mentorship",
            "scaling", "architecture", "infrastructure", "monitoring", "testing"
        ]

        found_keywords = []
        job_lower = job_description.lower()
        for keyword in common_keywords:
            if keyword.lower() in job_lower:
                found_keywords.append(keyword)

        return found_keywords

    # ===== CV DOCUMENT BUILDING HELPERS (Phase 8.2.2) =====

    def _extract_candidate_header(self, candidate_profile: str) -> Tuple[str, str]:
        """
        Extract candidate name and contact info from profile.

        Returns (name, contact_info_line) where contact_info_line is formatted as:
        "email | phone | linkedin"
        """
        lines = candidate_profile.split("\n")

        # Find name (usually first line or line after "Name:")
        name = "Candidate Name"
        for line in lines[:10]:
            if "Name:" in line:
                name = line.split("Name:")[1].strip()
                break
            elif line.strip() and not any(x in line for x in ["Email:", "Phone:", "LinkedIn:", "EXPERIENCE"]):
                # First non-empty, non-label line
                name = line.strip()
                break

        # Extract contact info
        contact_parts = []
        for line in lines[:15]:
            if "Email:" in line:
                email = line.split("Email:")[1].strip()
                contact_parts.append(email)
            elif "Phone:" in line:
                phone = line.split("Phone:")[1].strip()
                contact_parts.append(phone)
            elif "LinkedIn:" in line:
                linkedin = line.split("LinkedIn:")[1].strip()
                contact_parts.append(linkedin)

        contact_info = " | ".join(contact_parts) if contact_parts else ""

        return name, contact_info

    def _generate_professional_summary(
        self,
        job_title: str,
        company: str,
        fit_score: Optional[int],
        competency_mix: CompetencyMixOutput
    ) -> str:
        """
        Generate 2-3 sentence professional summary tailored to the job.

        Mentions fit_score if ≥80, and highlights top 2 competency dimensions.
        """
        # Get top 2 competencies
        numeric_competencies = {
            k: v for k, v in competency_mix.dict().items()
            if k != 'reasoning' and isinstance(v, (int, float))
        }
        top_competencies = sorted(numeric_competencies.items(), key=lambda x: x[1], reverse=True)[:2]

        summary_parts = []

        # Part 1: Role application (avoid mentioning target company to prevent hallucination QA false positives)
        summary_parts.append(
            f"Seasoned technical professional with demonstrated expertise in "
            f"{top_competencies[0][0]} and {top_competencies[1][0]}, targeting {job_title} opportunities."
        )

        # Part 2: Fit score if ≥80
        if fit_score and fit_score >= 80:
            summary_parts.append(
                f"Strong strategic alignment with this opportunity (fit score: {fit_score}/100)."
            )

        # Part 3: Value proposition
        summary_parts.append(
            "Proven track record of delivering measurable business impact through technical leadership and data-driven decision-making."
        )

        return " ".join(summary_parts)

    def _extract_key_achievements(self, selected_stars: List[Dict]) -> List[str]:
        """
        Extract 3-5 key achievement bullets from selected STARs.

        Each bullet starts with impact verb + includes quantified metric from results.
        Format: "Impact verb + action + metric"
        Example: "Reduced release cycle from 12 weeks to 2 weeks, improving team velocity by 6x"
        """
        achievements = []

        for star in selected_stars[:5]:  # Max 5 achievements
            # Extract results (should have metrics)
            results = star.get("results", "")
            metrics = star.get("metrics", "")

            # Build achievement bullet from results
            # If results already starts with an impact verb, use it directly
            if results:
                achievements.append(results)
            elif metrics:
                # Fallback: Use metrics if results is empty
                achievements.append(f"Achieved {metrics}")

        return achievements[:5]  # Ensure max 5

    def _format_star_as_bullets(self, star: Dict) -> List[str]:
        """
        Format a STAR record as 2-3 bullet points for experience section.

        Returns list of bullets covering situation/task, actions, and results.
        """
        bullets = []

        # Bullet 1: Context (situation + task)
        situation = star.get("situation", "")
        task = star.get("task", "")

        if situation and task:
            bullets.append(f"{situation.split('.')[0]}. Tasked with {task.lower()}")
        elif task:
            bullets.append(task)
        elif situation:
            bullets.append(situation)

        # Bullet 2: Actions
        actions = star.get("actions", "")
        if actions:
            bullets.append(actions)

        # Bullet 3: Results (with metrics)
        results = star.get("results", "")
        if results:
            bullets.append(results)

        return bullets

    def _extract_education(self, candidate_profile: str) -> str:
        """
        Extract education section from candidate profile.

        Returns formatted education string or minimal placeholder if not found.
        """
        # Look for education section
        lines = candidate_profile.split("\n")

        education_lines = []
        in_education = False

        for line in lines:
            if "EDUCATION" in line.upper() or "Education:" in line:
                in_education = True
                continue

            if in_education:
                # Stop at next major section
                if line.strip() and line.strip().isupper() and len(line.strip()) > 10:
                    break

                if line.strip():
                    education_lines.append(line.strip())

        if education_lines:
            return "\n".join(education_lines)
        else:
            # Return minimal placeholder that doesn't fabricate credentials
            return "[Education details available upon request]"

    def _generate_minimal_cv(self, state: JobState, competency_mix: CompetencyMixOutput) -> Tuple[str, str]:
        """
        Generate minimal CV when no STAR data is available (Phase 8.2.3).

        Returns (cv_path, cv_reasoning) tuple.

        Creates a basic CV from candidate_profile only, skipping STAR-based sections.
        """
        self.logger.info("Building minimal CV from profile...")

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # ===== HEADER =====
        candidate_name, contact_info = self._extract_candidate_header(state["candidate_profile"])

        name_para = doc.add_paragraph(candidate_name)
        name_para.runs[0].font.size = Pt(18)
        name_para.runs[0].font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if contact_info:
            contact_para = doc.add_paragraph(contact_info)
            if contact_para.runs:
                contact_para.runs[0].font.size = Pt(10)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # ===== PROFESSIONAL SUMMARY =====
        doc.add_heading("Professional Summary", level=2)

        summary = f"Applying for the {state['title']} role at {state['company']}. "\
                  f"Candidate profile available but detailed achievement metrics (STAR records) are not yet populated in the system."

        doc.add_paragraph(summary)

        # ===== PROFILE SUMMARY =====
        doc.add_heading("Background", level=2)

        # Extract a few lines from profile as placeholder
        profile_lines = state["candidate_profile"].split("\n")
        for line in profile_lines[5:15]:  # Skip header, take next 10 lines
            if line.strip() and not any(x in line for x in ["Name:", "Email:", "Phone:", "LinkedIn:"]):
                doc.add_paragraph(line.strip())

        # ===== EDUCATION =====
        doc.add_heading("Education & Certifications", level=2)
        education = self._extract_education(state["candidate_profile"])
        doc.add_paragraph(education)

        # Save CV
        cv_path = self._save_cv_document(state, doc)
        self.logger.info(f"Minimal CV saved: {cv_path}")

        # Generate cv_reasoning explaining limitation
        cv_reasoning = (
            f"Minimal CV generated for {state['title']} at {state['company']} due to missing STAR achievement data. "
            f"CV was derived solely from candidate profile text without competency scoring or achievement selection. "
            f"To enable STAR-driven tailoring, populate the STAR library with structured achievements."
        )

        return cv_path, cv_reasoning

    def _build_cv_document(self, state: JobState, competency_mix: CompetencyMixOutput, selected_stars: List[Dict]) -> Tuple[Document, str]:
        """
        Build CV document as .docx (Phase 8.2.2).

        Returns (doc, text_content) where:
        - doc: python-docx Document object ready to save
        - text_content: Plain text representation for hallucination QA
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Track text content for QA
        text_lines = []

        # ===== HEADER =====
        candidate_name, contact_info = self._extract_candidate_header(state["candidate_profile"])

        # Name
        name_para = doc.add_paragraph(candidate_name)
        name_para.runs[0].font.size = Pt(18)
        name_para.runs[0].font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_lines.append(candidate_name)

        # Contact info
        if contact_info:
            contact_para = doc.add_paragraph(contact_info)
            if contact_para.runs:
                contact_para.runs[0].font.size = Pt(10)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text_lines.append(contact_info)

        doc.add_paragraph()  # Blank line

        # ===== PROFESSIONAL SUMMARY =====
        doc.add_heading("Professional Summary", level=2)
        text_lines.append("Professional Summary")

        summary = self._generate_professional_summary(
            state["title"],
            state["company"],
            state.get("fit_score"),
            competency_mix
        )
        doc.add_paragraph(summary)
        text_lines.append(summary)

        # ===== KEY ACHIEVEMENTS =====
        doc.add_heading("Key Achievements", level=2)
        text_lines.append("Key Achievements")

        achievements = self._extract_key_achievements(selected_stars)
        for achievement in achievements:
            p = doc.add_paragraph(achievement, style='List Bullet')
            p.runs[0].font.size = Pt(11)
            text_lines.append(f"• {achievement}")

        # ===== PROFESSIONAL EXPERIENCE =====
        doc.add_heading("Professional Experience", level=2)
        text_lines.append("Professional Experience")

        # Sort STARs by period (reverse chronological)
        sorted_stars = sorted(selected_stars, key=lambda s: s.get('period', ''), reverse=True)

        for star in sorted_stars:
            # Role heading: "Engineering Manager | AdTech Co | 2020-2023"
            role_line = f"{star.get('role', 'Role')} | {star.get('company', 'Company')} | {star.get('period', 'Dates')}"
            role_para = doc.add_paragraph(role_line)
            role_para.runs[0].font.bold = True
            role_para.runs[0].font.size = Pt(11)
            text_lines.append(role_line)

            # STAR-based bullets
            star_bullets = self._format_star_as_bullets(star)
            for bullet in star_bullets:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.runs[0].font.size = Pt(10)
                text_lines.append(f"  • {bullet}")

            doc.add_paragraph()  # Blank line between roles

        # ===== EDUCATION & CERTIFICATIONS =====
        doc.add_heading("Education & Certifications", level=2)
        text_lines.append("Education & Certifications")

        education = self._extract_education(state["candidate_profile"])
        doc.add_paragraph(education)
        text_lines.append(education)

        return doc, "\n".join(text_lines)

    def _save_cv_document(self, state: JobState, doc: Document) -> str:
        """
        Save CV document as .docx (Phase 8.2.2).

        Args:
            state: JobState with company and title
            doc: python-docx Document object

        Returns:
            Path to saved .docx file
        """
        # Create output directory
        company_clean = state["company"].replace(" ", "_").replace("/", "_")
        title_clean = state["title"].replace(" ", "_").replace("/", "_")
        output_dir = Path("applications") / company_clean / title_clean
        output_dir.mkdir(parents=True, exist_ok=True)

        # Save as .docx
        cv_path = output_dir / f"CV_{company_clean}.docx"
        doc.save(str(cv_path))

        return str(cv_path)
