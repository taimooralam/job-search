"""
Layer 6: Outreach & CV Generator

Generates personalized cover letter and tailored CV based on all previous analysis.
This is the SIMPLIFIED version for today's vertical slice.

FUTURE: Will expand to include per-person outreach templates, richer CV formatting.
"""

import os
import tempfile
from datetime import datetime
from typing import Dict, Any, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from tenacity import retry, stop_after_attempt, wait_exponential

from src.common.config import Config
from src.common.state import JobState


# ===== OUTREACH GENERATION PROMPTS =====

OUTREACH_SYSTEM_PROMPT = """You are an expert career consultant and professional writer specializing in compelling cover letters.

Your task: Write a concise, professional 3-paragraph cover letter that demonstrates strong fit for the role.

Structure:
- Paragraph 1: Express interest in the specific role and company, mentioning a key company detail
- Paragraph 2: Highlight 2-3 most relevant experiences that address the job's pain points
- Paragraph 3: Express enthusiasm and call to action

Style:
- Professional but warm
- Specific and evidence-based (use concrete achievements)
- Confident without being arrogant
- 250-350 words total

Do NOT include:
- Address/date header (we'll add that separately)
- Signature block
- Generic platitudes
"""

OUTREACH_USER_PROMPT_TEMPLATE = """Write a 3-paragraph cover letter for this opportunity:

=== JOB DETAILS ===
Title: {title}
Company: {company}
Company Summary: {company_summary}

=== PAIN POINTS (What they need) ===
{pain_points}

=== CANDIDATE PROFILE ===
{candidate_profile}

=== FIT ANALYSIS ===
Score: {fit_score}/100
Rationale: {fit_rationale}

=== YOUR TASK ===
Write a compelling 3-paragraph cover letter that:
1. Shows genuine interest in {company} and the {title} role
2. Highlights specific achievements that address their pain points
3. Demonstrates why this is a strong mutual fit

Cover Letter:
"""


class OutreachGenerator:
    """Generates personalized cover letter."""

    def __init__(self):
        """Initialize LLM for cover letter generation."""
        self.llm = ChatOpenAI(
            model=Config.DEFAULT_MODEL,
            temperature=Config.CREATIVE_TEMPERATURE,  # 0.7 for creative writing
            api_key=Config.get_llm_api_key(),
            base_url=Config.get_llm_base_url(),
        )

    def _format_pain_points(self, pain_points: list) -> str:
        """Format pain points as numbered list."""
        if not pain_points:
            return "No specific pain points identified."
        return "\n".join(f"{i}. {point}" for i, point in enumerate(pain_points, 1))

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        reraise=True
    )
    def generate_cover_letter(self, state: JobState) -> str:
        """
        Generate cover letter using LLM.

        Args:
            state: JobState with all previous layer outputs

        Returns:
            Cover letter text (3 paragraphs)
        """
        pain_points_text = self._format_pain_points(state.get("pain_points", []))

        messages = [
            SystemMessage(content=OUTREACH_SYSTEM_PROMPT),
            HumanMessage(
                content=OUTREACH_USER_PROMPT_TEMPLATE.format(
                    title=state["title"],
                    company=state["company"],
                    company_summary=state.get("company_summary", "No company information available."),
                    pain_points=pain_points_text,
                    candidate_profile=state.get("candidate_profile", "No candidate profile provided."),
                    fit_score=state.get("fit_score", "N/A"),
                    fit_rationale=state.get("fit_rationale", "No fit analysis available.")
                )
            )
        ]

        response = self.llm.invoke(messages)
        cover_letter = response.content.strip()

        return cover_letter


class CVGenerator:
    """Generates tailored CV document."""

    def __init__(self):
        """Initialize CV generator."""
        pass

    def _parse_candidate_profile(self, profile: str) -> Dict[str, Any]:
        """
        Parse candidate profile text into structured data.

        Returns dict with: name, email, phone, linkedin, experience, skills, achievements, work_history, education
        """
        # Simple parsing - extract key sections
        data = {
            "name": "Candidate Name",
            "email": "",
            "phone": "",
            "linkedin": "",
            "experience": "",
            "skills": "",
            "achievements": [],
            "work_history": [],
            "education": []
        }

        # Extract name (first line after "Name:")
        if "Name:" in profile:
            name_line = profile.split("Name:")[1].split("\n")[0].strip()
            data["name"] = name_line

        # Extract email
        if "Email:" in profile:
            email_line = profile.split("Email:")[1].split("\n")[0].strip()
            data["email"] = email_line

        # Extract phone
        if "Phone:" in profile:
            phone_line = profile.split("Phone:")[1].split("\n")[0].strip()
            data["phone"] = phone_line

        # Extract LinkedIn
        if "LinkedIn:" in profile:
            linkedin_line = profile.split("LinkedIn:")[1].split("\n")[0].strip()
            data["linkedin"] = linkedin_line

        # Extract experience summary
        if "Experience:" in profile:
            exp_line = profile.split("Experience:")[1].split("\n")[0].strip()
            data["experience"] = exp_line

        # Extract skills
        if "Key Skills:" in profile:
            skills_line = profile.split("Key Skills:")[1].split("\n")[0].strip()
            data["skills"] = skills_line

        # Extract achievements (lines starting with -)
        achievements = []
        if "Notable Achievements:" in profile or "Achievements:" in profile:
            section = profile.split("Achievements:")[1] if "Achievements:" in profile else ""
            if "Notable Achievements:" in profile:
                section = profile.split("Notable Achievements:")[1]

            for line in section.split("\n"):
                line = line.strip()
                if line.startswith("-"):
                    achievements.append(line[1:].strip())
        data["achievements"] = achievements[:4]  # Limit to top 4

        return data

    def generate_cv(self, state: JobState) -> str:
        """
        Generate tailored CV document.

        Args:
            state: JobState with candidate profile and job details

        Returns:
            Path to generated .docx file
        """
        # Parse candidate profile
        candidate = self._parse_candidate_profile(state.get("candidate_profile", ""))

        # Create document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # === HEADER ===
        # Name
        name_para = doc.add_paragraph(candidate["name"])
        name_para.runs[0].font.size = Pt(18)
        name_para.runs[0].font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = []
        if candidate["email"]:
            contact_parts.append(candidate["email"])
        if candidate["phone"]:
            contact_parts.append(candidate["phone"])
        if candidate["linkedin"]:
            contact_parts.append(candidate["linkedin"])

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            if contact_para.runs:
                contact_para.runs[0].font.size = Pt(10)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Blank line

        # === JOB-SPECIFIC SUMMARY ===
        doc.add_heading("Professional Summary", level=2)

        # Custom summary paragraph highlighting fit for THIS job
        summary_text = (
            f"Experienced marketing professional with {candidate.get('experience', '8+ years experience')} "
            f"applying for the {state['title']} role at {state['company']}. "
            f"Proven track record in {candidate.get('skills', 'performance marketing, data analysis, and strategic planning')}. "
        )

        if state.get("fit_score") and state.get("fit_score") >= 80:
            summary_text += f"Strong alignment with this opportunity (fit score: {state['fit_score']}/100)."

        doc.add_paragraph(summary_text)

        # === CORE COMPETENCIES ===
        if candidate["skills"]:
            doc.add_heading("Core Competencies", level=2)
            doc.add_paragraph(candidate["skills"])

        # === KEY ACHIEVEMENTS ===
        if candidate["achievements"]:
            doc.add_heading("Key Achievements", level=2)
            for achievement in candidate["achievements"]:
                p = doc.add_paragraph(achievement, style='List Bullet')
                p.runs[0].font.size = Pt(11)

        # === EXPERIENCE SUMMARY ===
        doc.add_heading("Professional Experience", level=2)
        doc.add_paragraph(
            f"Detailed work history available upon request. "
            f"Highlights include leadership in performance marketing, "
            f"data-driven campaign optimization, and team management across multiple organizations."
        )

        # === EDUCATION ===
        doc.add_heading("Education", level=2)
        doc.add_paragraph("MBA, Marketing\nBS, Business Administration")

        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_clean = state["company"].replace(" ", "_").replace("/", "_")
        filename = f"CV_{company_clean}_{timestamp}.docx"

        # Save to temp directory
        output_path = os.path.join(tempfile.gettempdir(), filename)
        doc.save(output_path)

        return output_path


class Generator:
    """Main class orchestrating cover letter and CV generation."""

    def __init__(self):
        """Initialize both generators."""
        self.outreach_gen = OutreachGenerator()
        self.cv_gen = CVGenerator()

    def generate_outputs(self, state: JobState) -> Dict[str, Any]:
        """
        Generate both cover letter and CV.

        Args:
            state: JobState with all previous layer outputs

        Returns:
            Dict with cover_letter and cv_path keys
        """
        try:
            print(f"   Generating cover letter...")
            cover_letter = self.outreach_gen.generate_cover_letter(state)
            print(f"   ✓ Cover letter generated ({len(cover_letter)} chars)")

            print(f"   Generating tailored CV...")
            cv_path = self.cv_gen.generate_cv(state)
            print(f"   ✓ CV generated: {cv_path}")

            return {
                "cover_letter": cover_letter,
                "cv_path": cv_path
            }

        except Exception as e:
            error_msg = f"Layer 6 (Generator) failed: {str(e)}"
            print(f"   ✗ {error_msg}")

            errors_list = state.get("errors") or []
            if isinstance(errors_list, str):
                errors_list = [errors_list]

            return {
                "cover_letter": None,
                "cv_path": None,
                "errors": errors_list + [error_msg]
            }


# ===== LANGGRAPH NODE FUNCTION =====

def generator_node(state: JobState) -> Dict[str, Any]:
    """
    LangGraph node function for Layer 6: Outreach & CV Generator.

    Args:
        state: Current job processing state

    Returns:
        Dictionary with updates to merge into state
    """
    print("\n" + "="*60)
    print("LAYER 6: Outreach & CV Generator")
    print("="*60)

    generator = Generator()
    updates = generator.generate_outputs(state)

    # Print results
    if updates.get("cover_letter"):
        print(f"\n📄 Cover Letter Preview (first 150 chars):")
        print(f"  {updates['cover_letter'][:150]}...")
    else:
        print("\n⚠️  No cover letter generated")

    if updates.get("cv_path"):
        print(f"\n📋 CV Generated: {updates['cv_path']}")
    else:
        print("\n⚠️  No CV generated")

    print("="*60 + "\n")

    return updates
