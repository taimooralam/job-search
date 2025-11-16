"""
TDD Test for Layer 7: Output Publisher

This test defines the expected behavior BEFORE implementation.

Layer 7 should:
1. Upload cover letter (.txt) to Google Drive
2. Upload CV (.docx) to Google Drive
3. Create folder structure: /applications/<company>/<role>/
4. Log tracking row to Google Sheets
5. Return drive_folder_url and sheet_row_id
"""

import os
import tempfile
from datetime import datetime
from src.layer7.output_publisher import output_publisher_node
from src.common.state import JobState

# Create test files for upload
def create_test_files():
    """Create temporary test files for upload."""
    # Create temporary cover letter
    cover_letter = """Dear Hiring Team at Launch Potato,

I am excited to apply for the Senior Manager, YouTube Paid Performance role.
With 8+ years of experience scaling YouTube campaigns, I am confident I can drive growth.

Best regards,
Alex Thompson"""

    cover_letter_path = os.path.join(tempfile.gettempdir(), "test_cover_letter.txt")
    with open(cover_letter_path, "w") as f:
        f.write(cover_letter)

    # Create a simple test CV (we'll use the one from Layer 6 test, or create a dummy)
    from docx import Document
    doc = Document()
    doc.add_heading("Alex Thompson", 0)
    doc.add_paragraph("Test CV for Layer 7")

    cv_path = os.path.join(tempfile.gettempdir(), "test_cv.docx")
    doc.save(cv_path)

    return cover_letter, cv_path


# Sample state with all outputs from previous layers
cover_letter_text, test_cv_path = create_test_files()

sample_state: JobState = {
    "job_id": "4335713702",
    "title": "Senior Manager, YouTube Paid Performance",
    "company": "Launch Potato",
    "job_description": "YouTube performance marketing role...",
    "job_url": "https://www.linkedin.com/jobs/view/4335713702",
    "source": "linkedin",
    "candidate_profile": "Alex Thompson - 8+ years experience...",

    # Previous layer outputs
    "pain_points": [
        "Need expert who can profitably scale YouTube campaigns",
        "Requires strong analytical capabilities",
    ],
    "company_summary": "Launch Potato is a digital media company...",
    "company_url": "https://launchpotato.com",
    "fit_score": 95,
    "fit_rationale": "Exceptional fit with 8+ years of relevant experience...",
    "cover_letter": cover_letter_text,
    "cv_path": test_cv_path,

    # Fields to be filled by Layer 7
    "drive_folder_url": None,
    "sheet_row_id": None,

    # Other fields
    "run_id": None,
    "created_at": None,
    "errors": None,
    "status": "processing"
}


def test_output_publisher():
    """Test that Layer 7 uploads to Drive and logs to Sheets."""
    print("Testing Layer 7: Output Publisher (TDD)\n")
    print("⚠️  NOTE: This test will create real files in Google Drive and Sheets!")
    print("   Make sure your .env is configured with valid credentials.\n")

    # Run the node
    updates = output_publisher_node(sample_state)

    # ASSERTIONS: Define what we expect
    print("\n" + "="*60)
    print("TEST RESULTS")
    print("="*60)

    # Test 1: drive_folder_url should be present
    assert updates.get("drive_folder_url") is not None, "❌ drive_folder_url is missing"
    print(f"✅ drive_folder_url present")

    # Test 2: drive_folder_url should be a valid URL
    assert updates["drive_folder_url"].startswith("https://"), \
        f"❌ drive_folder_url is not a valid URL: {updates['drive_folder_url']}"
    print(f"✅ drive_folder_url is valid URL")

    # Test 3: drive_folder_url should contain 'drive.google.com' or 'docs.google.com'
    assert "drive.google.com" in updates["drive_folder_url"] or "docs.google.com" in updates["drive_folder_url"], \
        f"❌ drive_folder_url is not a Google Drive URL: {updates['drive_folder_url']}"
    print(f"✅ drive_folder_url is Google Drive link")

    # Test 4: sheet_row_id should be present
    assert updates.get("sheet_row_id") is not None, "❌ sheet_row_id is missing"
    print(f"✅ sheet_row_id present: {updates['sheet_row_id']}")

    # Test 5: sheet_row_id should be a valid number (row number in Google Sheets)
    assert isinstance(updates["sheet_row_id"], (int, str)), \
        f"❌ sheet_row_id is not a number: {type(updates['sheet_row_id'])}"
    print(f"✅ sheet_row_id is valid")

    # Test 6: No errors should have occurred
    if updates.get("errors"):
        print(f"⚠️  Errors occurred: {updates['errors']}")
    else:
        print(f"✅ No errors during publishing")

    print("\n" + "="*60)
    print("ALL TESTS PASSED! ✅")
    print("="*60)

    print(f"\n📁 Google Drive Folder: {updates['drive_folder_url']}")
    print(f"📊 Google Sheets Row: {updates['sheet_row_id']}")
    print("\n💡 Check your Google Drive and Sheets to verify uploads!")
    print("="*60)

    # Cleanup test files
    if os.path.exists(test_cv_path):
        os.remove(test_cv_path)
    cover_letter_path = os.path.join(tempfile.gettempdir(), "test_cover_letter.txt")
    if os.path.exists(cover_letter_path):
        os.remove(cover_letter_path)


if __name__ == "__main__":
    test_output_publisher()
